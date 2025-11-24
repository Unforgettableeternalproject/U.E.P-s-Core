import os
import platform
import subprocess
import shutil as shutil_module
from pathlib import Path
from utils.debug_helper import debug_log, debug_log_e, info_log, error_log


def drop_and_read(file_path: str) -> str:
    """
    將拖放的檔案內容讀取為文字 - 支援 txt、md、pdf 格式
    
    Args:
        file_path: 要讀取的檔案路徑
        
    Returns:
        檔案的文字內容
        
    Raises:
        ValueError: 當檔案格式不支援時
        Exception: 其他讀取錯誤
    """
    info_log(f"[file] 讀取檔案：{file_path}")
    ext = Path(file_path).suffix.lower()
    try:
        if ext in (".txt", ".md"):
            return Path(file_path).read_text(encoding="utf-8")
        elif ext == ".pdf":
            from pdfminer.high_level import extract_text # type: ignore
            return extract_text(file_path)
        else:
            raise ValueError(f"不支援格式：{ext}")
    except Exception as e:
        error_log(f"[file] 讀取失敗：{e}")
        raise

def intelligent_archive(file_path: str, target_dir: str = "") -> str:
    """
    智能歸檔功能 - 根據檔案類型、使用者指定或歷史記錄智能地將檔案歸檔到適當位置
    
    Args:
        file_path: 要歸檔的檔案路徑
        target_dir: 使用者指定的目標資料夾，若為空則自動判斷
        
    Returns:
        歸檔後的新檔案路徑
    """
    import shutil
    import json
    import datetime
    from pathlib import Path
    
    info_log(f"[file] 準備歸檔檔案：{file_path}")
    
    # 檢查檔案是否存在
    file_path_obj = Path(file_path)
    if not file_path_obj.exists():
        error_log(f"[file] 歸檔失敗：檔案 {file_path} 不存在")
        raise FileNotFoundError(f"檔案 {file_path} 不存在")
    
    # 取得檔案資訊
    file_name = file_path_obj.name
    file_ext = file_path_obj.suffix.lower()
    
    # 如果使用者有指定目標資料夾，優先使用
    if target_dir and Path(target_dir).exists():
        target_path = Path(target_dir) / file_name
        info_log(f"[file] 使用用戶指定位置：{target_dir}")
    else:
        # 查看過去的歸檔記錄
        archive_history_path = Path("configs") / "archive_history.json"
        similar_target = None
        
        # 檢查歷史記錄
        if archive_history_path.exists():
            try:
                with open(archive_history_path, "r", encoding="utf-8") as f:
                    history = json.load(f)
                
                # 尋找相似檔案的歸檔位置
                for record in history.get("archives", []):
                    if record.get("file_ext") == file_ext:
                        similar_target = record.get("target_dir")
                        break
                        
                if similar_target:
                    info_log(f"[file] 找到相似檔案歷史記錄，使用位置：{similar_target}")
            except Exception as e:
                error_log(f"[file] 讀取歸檔記錄失敗：{e}")
        
        # 如果找到了相似的歸檔記錄
        if similar_target and Path(similar_target).exists():
            target_path = Path(similar_target) / file_name
        else:
            # 根據檔案類型選擇預設位置
            if file_ext in (".txt", ".doc", ".docx", ".pdf", ".md", ".csv", ".xlsx", ".pptx"):
                # 文件類型歸到文件資料夾
                target_path = Path(os.path.expanduser("~/Documents")) / file_name
            elif file_ext in (".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff"):
                # 圖片類型歸到圖片資料夾
                target_path = Path(os.path.expanduser("~/Pictures")) / file_name
            elif file_ext in (".mp3", ".wav", ".flac", ".ogg", ".m4a"):
                # 音樂類型歸到音樂資料夾
                target_path = Path(os.path.expanduser("~/Music")) / file_name
            elif file_ext in (".mp4", ".avi", ".mkv", ".mov", ".wmv"):
                # 影片類型歸到影片資料夾
                target_path = Path(os.path.expanduser("~/Videos")) / file_name
            else:
                # 其他類型歸到下載資料夾
                target_path = Path(os.path.expanduser("~/Downloads")) / file_name
            
            info_log(f"[file] 根據檔案類型選擇預設位置：{target_path}")
    
    # 確保目標資料夾存在
    target_path.parent.mkdir(parents=True, exist_ok=True)
    
    # 處理同名檔案
    if target_path.exists():
        base = target_path.stem
        extension = target_path.suffix
        counter = 1
        
        while True:
            new_name = f"{base}_{counter}{extension}"
            new_target = target_path.parent / new_name
            if not new_target.exists():
                target_path = new_target
                break
            counter += 1
    
    # 移動檔案
    try:
        shutil.copy2(file_path, target_path)
        info_log(f"[file] 檔案已複製至：{target_path}")
        
        # 記錄歸檔資訊
        archive_history_path = Path("configs") / "archive_history.json"
        
        history = {"archives": []}
        if archive_history_path.exists():
            try:
                with open(archive_history_path, "r", encoding="utf-8") as f:
                    history = json.load(f)
            except:
                pass
        
        # 添加新記錄
        new_record = {
            "original_path": str(file_path),
            "target_dir": str(target_path.parent),
            "file_name": file_name,
            "file_ext": file_ext,
            "timestamp": datetime.datetime.now().isoformat()
        }
        
        if "archives" not in history:
            history["archives"] = []
        
        history["archives"].insert(0, new_record)  # 最新記錄放在最前面
        
        # 限制記錄數量，避免檔案過大
        history["archives"] = history["archives"][:100]
        
        # 保存記錄
        try:
            with open(archive_history_path, "w", encoding="utf-8") as f:
                json.dump(history, f, ensure_ascii=False, indent=2)
        except Exception as e:
            error_log(f"[file] 保存歸檔記錄失敗：{e}")
        
        return str(target_path)
    except Exception as e:
        error_log(f"[file] 歸檔失敗：{e}")
        raise

def summarize_tag(file_path: str, tag_count: int = 3) -> dict:
    """
    為檔案生成摘要與標籤 - 解析檔案內容，使用LLM模組生成摘要，輸出至同路徑的summary檔案
    
    Args:
        file_path: 要摘要的檔案路徑
        tag_count: 要生成的標籤數量，默認為3個
        
    Returns:
        包含摘要檔案路徑的字典
    """
    import os
    import json
    import datetime
    import importlib
    from pathlib import Path
    
    info_log(f"[file] 準備摘要並標記檔案：{file_path}")
    
    # 檢查檔案是否存在
    file_path_obj = Path(file_path)
    if not file_path_obj.exists():
        error_log(f"[file] 摘要失敗：檔案 {file_path} 不存在")
        raise FileNotFoundError(f"檔案 {file_path} 不存在")
    
    # 讀取檔案內容
    try:
        content = drop_and_read(file_path)  # 使用已實作的函數讀取檔案
    except Exception as e:
        error_log(f"[file] 摘要失敗，無法讀取檔案：{e}")
        raise
    
    # 準備生成摘要
    summary_content = ""
    tags = []
    
    # 嘗試使用LLM模組生成摘要，如果模組未啟用則使用簡單摘要
    try:
        # 嘗試導入並使用LLM模組
        try:            # 動態導入LLM模組
            llm_module = None
            try:
                from modules.llm_module.llm_module import LLMModule
                from configs.config_loader import load_module_config
                
                config = load_module_config("llm_module")
                # 禁用隱性快取，避免測試影響系統快取
                if "use_prompt_caching" in config:
                    config["use_prompt_caching"] = False
                llm_module = LLMModule(config)
                info_log(f"[file] 成功載入LLM模組（測試模式，已禁用快取）")
            except ImportError as e:
                info_log(f"[file] LLM模組未啟用或無法導入，將使用簡單摘要: {e}")
                
            if llm_module:
                # 構建清晰的摘要提示詞
                prompt = f"""Summarize the file and generate tags:

File content：
{content[:5000]}{"..." if len(content) > 5000 else ""}

Please respond in the following format:
Tags: Tag1, Tag2, Tag3{"" if tag_count == 3 else f", Tag{tag_count}"}
Summary: [Write the summary here]

Requirements:
1. Generate {tag_count} relevant key tags
2. Provide a concise yet comprehensive summary
3. Tags should reflect the main themes and content features of the file
"""
                
                # 構建摘要請求 (需要符合LLMInput格式)
                request_data = {
                    "text": prompt,
                    "intent": "chat",  # LLMInput需要，目前僅支援chat
                    "is_internal": True  # 使用內部調用模式，避免加入系統指示詞
                }
                
                # 呼叫LLM模組
                response = llm_module.handle(request_data)
                
                if response and "text" in response and response.get("status") == "ok":
                    # 從LLM回應的text中提取摘要和標籤
                    llm_response_text = response["text"]
                    info_log(f"[file] LLM 回應: {llm_response_text[:200]}...")
                    
                    # 改進的標籤和摘要解析邏輯
                    try:
                        # 提取標籤部分
                        if "Tags:" in llm_response_text or "Tags：" in llm_response_text:
                            # 查找標籤行
                            lines = llm_response_text.split('\n')
                            tags_line = ""
                            summary_lines = []
                            found_tags = False
                            found_summary = False
                            
                            for line in lines:
                                line = line.strip()
                                if not found_tags and ("Tags:" in line or "Tags：" in line):
                                    tags_line = line.split("：")[1] if "：" in line else line.split(":")[1]
                                    found_tags = True
                                elif found_tags and ("Summary:" in line or "Summary：" in line):
                                    summary_start = line.split("：")[1] if "：" in line else line.split(":")[1]
                                    if summary_start.strip():
                                        summary_lines.append(summary_start.strip())
                                    found_summary = True
                                elif found_summary:
                                    if line:
                                        summary_lines.append(line)
                            
                            # 解析標籤
                            if tags_line:
                                tags = [tag.strip() for tag in tags_line.split(',')]
                                # 清理標籤，移除空字符串並限制數量
                                tags = [tag for tag in tags if tag][:tag_count]
                            else:
                                tags = [file_path_obj.stem]  # fallback
                            
                            # 組合摘要
                            if summary_lines:
                                summary_content = '\n'.join(summary_lines).strip()
                            else:
                                summary_content = llm_response_text  # 使用整個回應作為摘要
                        else:
                            # 如果沒有找到格式化的回應，嘗試智能解析
                            info_log("[file] LLM 回應格式不規範，嘗試智能解析")
                            # 使用整個回應作為摘要
                            summary_content = llm_response_text
                            # 基於檔案生成簡單標籤
                            tags = [file_path_obj.stem, file_path_obj.suffix.lower()[1:] if file_path_obj.suffix else "file"]
                            if tag_count > 2:
                                tags.append("document")
                        
                        info_log(f"[file] LLM模組成功生成摘要和{len(tags)}個標籤: {tags}")
                        
                    except Exception as parse_error:
                        error_log(f"[file] 解析 LLM 回應失敗: {parse_error}")
                        # fallback 到整個回應
                        summary_content = llm_response_text
                        tags = [file_path_obj.stem]
                else:
                    raise ValueError(f"LLM模組未返回有效摘要: {response.get('status', 'unknown error') if response else 'no response'}")
        except Exception as e:
            info_log(f"[file] 使用LLM模組摘要失敗：{e}，將使用簡單摘要")
            # 若LLM模組失敗，使用簡單摘要法
            raise
            
        # 如果沒有成功使用LLM模組，使用簡單方法生成摘要
        if not summary_content:
            raise ValueError("需要使用簡單摘要")
            
    except Exception as e:
        # 使用簡單摘要方法
        info_log(f"[file] 使用簡單摘要方法")
        
        # 簡單摘要：取前1000個字符
        summary_preview = content[:1000] + ("..." if len(content) > 1000 else "")
        
        # 簡單標籤：使用檔案類型和大小
        file_ext = file_path_obj.suffix.lower()[1:]  # 移除開頭的點
        file_size_kb = os.path.getsize(file_path) / 1024
        file_size_tag = f"{file_size_kb:.1f}KB" if file_size_kb < 1024 else f"{file_size_kb/1024:.1f}MB"
        
        tags = [file_ext, file_size_tag, file_path_obj.name.split(".")[0]]
        summary_content = f"檔案預覽：\n\n{summary_preview}\n\n(自動生成的簡單摘要，內容為檔案前1000個字符)"
    
    # 生成摘要檔案的內容
    output_content = f"""# {file_path_obj.name} 摘要

## 檔案資訊
- 原始檔案：{file_path_obj.name}
- 建立時間：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

## 標籤
{', '.join(f'`{tag}`' for tag in tags)}

## 內容摘要
{summary_content}
"""
    
    # 建立摘要檔案
    summary_file_path = file_path_obj.parent / f"{file_path_obj.stem}-summary.md"
    try:
        with open(summary_file_path, "w", encoding="utf-8") as f:
            f.write(output_content)
        info_log(f"[file] 成功生成摘要檔案：{summary_file_path}")
    except Exception as e:
        error_log(f"[file] 建立摘要檔案失敗：{e}")
        raise
    
    # 返回摘要檔案資訊
    return {
        "summary_file": str(summary_file_path),
        "tags": tags
    }


def clean_trash_bin() -> str:
    """
    清空系統資源回收桶 - 支援 Windows、macOS 和 Linux 跨平台
    
    Returns:
        清空結果訊息
        
    Raises:
        Exception: 清空失敗時拋出異常
    """
    system = platform.system()
    info_log(f"[file] 準備清空資源回收桶（系統：{system}）")
    
    try:
        if system == "Windows":
            result = subprocess.run(
                ["powershell", "-Command", "Clear-RecycleBin -Force"],
                capture_output=True, text=True
            )
            if result.returncode == 0 or result.returncode == 1:
                # returncode 1 在某些情況下是正常的（如回收桶已空）
                info_log("[file] Windows 資源回收桶已清空")
                return "Windows 資源回收桶已成功清空"
            else:
                error_log(f"[file] 清空失敗，錯誤碼：{result.returncode}，輸出：{result.stderr}")
                raise Exception(f"清空資源回收桶失敗：{result.stderr}")

        elif system == "Darwin":  # macOS
            trash_path = os.path.expanduser("~/.Trash")
            if os.path.exists(trash_path):
                for file in os.listdir(trash_path):
                    file_path = os.path.join(trash_path, file)
                    if os.path.isfile(file_path) or os.path.islink(file_path):
                        os.remove(file_path)
                    elif os.path.isdir(file_path):
                        shutil_module.rmtree(file_path)
                info_log("[file] macOS 資源回收桶已清空")
                return "macOS 資源回收桶已成功清空"
            else:
                info_log("[file] macOS 資源回收桶路徑不存在或已空")
                return "macOS 資源回收桶已空"

        elif system == "Linux":
            trash_paths = [
                os.path.expanduser("~/.local/share/Trash/files"),
                os.path.expanduser("~/.local/share/Trash/info")
            ]
            for path in trash_paths:
                if os.path.exists(path):
                    for file in os.listdir(path):
                        file_path = os.path.join(path, file)
                        if os.path.isfile(file_path) or os.path.islink(file_path):
                            os.remove(file_path)
                        elif os.path.isdir(file_path):
                            shutil_module.rmtree(file_path)
            info_log("[file] Linux 資源回收桶已清空")
            return "Linux 資源回收桶已成功清空"

        else:
            error_log(f"[file] 不支援的作業系統：{system}")
            raise Exception(f"不支援的作業系統：{system}")

    except Exception as e:
        error_log(f"[file] 清空資源回收桶失敗：{e}")
        raise


def translate_document(
    file_path: str, 
    target_lang: str = "zh-tw", 
    source_lang: str = "auto",
    output_path: str = ""
) -> str:
    """
    翻譯文件內容並匯出 - 支援 PDF、DOCX、TXT 多格式
    
    Args:
        file_path: 要翻譯的檔案路徑
        target_lang: 目標語言（預設繁體中文 zh-tw）
        source_lang: 來源語言（預設自動偵測 auto）
        output_path: 輸出檔案路徑（若為空則自動生成）
        
    Returns:
        翻譯後的檔案路徑
        
    Raises:
        ValueError: 不支援的檔案格式
        Exception: 翻譯或儲存失敗
    """
    info_log(f"[file] 準備翻譯文件：{file_path}")
    
    file_path_obj = Path(file_path)
    if not file_path_obj.exists():
        error_log(f"[file] 翻譯失敗：檔案 {file_path} 不存在")
        raise FileNotFoundError(f"檔案 {file_path} 不存在")
    
    file_ext = file_path_obj.suffix.lower()
    
    # 偵測檔案類型並提取文本
    try:
        text_content = _extract_text_from_file(file_path, file_ext)
        if not text_content:
            raise ValueError("文件內容為空")
        info_log(f"[file] 成功提取文本，長度：{len(text_content)} 字元")
    except Exception as e:
        error_log(f"[file] 提取文本失敗：{e}")
        raise
    
    # 翻譯文本
    try:
        translated_text = _translate_text(text_content, target_lang, source_lang)
        info_log(f"[file] 翻譯完成，長度：{len(translated_text)} 字元")
    except Exception as e:
        error_log(f"[file] 翻譯失敗：{e}")
        raise
    
    # 儲存翻譯結果
    try:
        if not output_path:
            # 自動生成輸出路徑（同目錄，加上 _translated 後綴）
            output_path = str(file_path_obj.parent / f"{file_path_obj.stem}_translated{file_ext}")
        
        _save_translated_file(translated_text, output_path, file_ext)
        info_log(f"[file] 翻譯檔案已儲存至：{output_path}")
        return output_path
    except Exception as e:
        error_log(f"[file] 儲存翻譯檔案失敗：{e}")
        raise


def _extract_text_from_file(file_path: str, file_ext: str) -> str:
    """從不同格式的檔案中提取文本"""
    if file_ext == ".txt" or file_ext == ".md":
        return Path(file_path).read_text(encoding="utf-8")
    
    elif file_ext == ".pdf":
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            return "\n\n".join(text_parts)
        except ImportError:
            error_log("[file] pdfplumber 未安裝，請執行：pip install pdfplumber")
            raise ImportError("需要安裝 pdfplumber：pip install pdfplumber")
    
    elif file_ext == ".docx":
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n\n".join(paragraphs)
        except ImportError:
            error_log("[file] python-docx 未安裝，請執行：pip install python-docx")
            raise ImportError("需要安裝 python-docx：pip install python-docx")
    
    else:
        raise ValueError(f"不支援的檔案格式：{file_ext}")


def _translate_text(text: str, target_lang: str, source_lang: str) -> str:
    """翻譯文本（使用 LLM）"""
    try:
        from modules.llm_module.llm_module import LLMModule
        from configs.config_loader import load_module_config
        from .maps.language_map import lang_map
        
        # 初始化 LLM 模組
        config = load_module_config("llm_module")
        if "use_prompt_caching" in config:
            config["use_prompt_caching"] = False  # 禁用快取避免測試影響
        llm_module = LLMModule(config)
        
        # 建立反向映射（代碼 -> 名稱）和擴展的語言名稱
        code_to_name = {code: name for name, code in lang_map.items()}
        lang_names = {
            "zh-tw": "Traditional Chinese (繁體中文)",
            "zh-cn": "Simplified Chinese (简体中文)",
            "en": "English (英文)",
            "ja": "Japanese (日本語)",
            "ko": "Korean (한국어)",
            "fr": "French (法文)",
            "de": "German (德文)",
            "es": "Spanish (西班牙文)",
            "pt": "Portuguese (葡萄牙文)",
            "it": "Italian (義大利文)",
            "ru": "Russian (俄文)",
            "ar": "Arabic (阿拉伯文)",
            "th": "Thai (泰文)",
            "vi": "Vietnamese (越南文)",
            "hi": "Hindi (印地文)",
            "id": "Indonesian (印尼文)",
            "nl": "Dutch (荷蘭文)",
            "el": "Greek (希臘文)",
            "tr": "Turkish (土耳其文)",
            "sv": "Swedish (瑞典文)",
            "auto": "auto-detect"
        }
        
        # 支援中文名稱或代碼輸入
        if target_lang in lang_map:
            target_lang = lang_map[target_lang]  # 中文名稱轉代碼
        if source_lang in lang_map:
            source_lang = lang_map[source_lang]  # 中文名稱轉代碼
            
        target_lang_name = lang_names.get(target_lang.lower(), target_lang)
        source_lang_name = lang_names.get(source_lang.lower(), "auto-detect")
        
        # 分段處理長文本（避免 token 限制）
        segments = _segment_text(text, max_length=8000)  # LLM 可以處理更長的文本
        translated_segments = []
        
        info_log(f"[file] 使用 LLM 翻譯，共 {len(segments)} 個段落")
        
        for i, segment in enumerate(segments):
            try:
                # 構建翻譯提示詞
                if source_lang == "auto":
                    prompt = f"""Please translate the following text to {target_lang_name}.
Maintain the original formatting and structure.

Text to translate:
{segment}

Translated text:"""
                else:
                    prompt = f"""Please translate the following text from {source_lang_name} to {target_lang_name}.
Maintain the original formatting and structure.

Text to translate:
{segment}

Translated text:"""
                
                # 使用內部呼叫模式（繞過會話檢查）
                request_data = {
                    "text": prompt,
                    "intent": "chat",
                    "is_internal": True
                }
                
                response = llm_module.handle(request_data)
                
                if response and response.get("status") == "ok" and "text" in response:
                    translated_segments.append(response["text"].strip())
                    info_log(f"[file] 翻譯進度：{i+1}/{len(segments)}")
                else:
                    error_log(f"[file] 翻譯段落 {i+1} 失敗，保留原文")
                    translated_segments.append(segment)
            
            except Exception as seg_error:
                error_log(f"[file] 翻譯段落 {i+1} 失敗：{seg_error}")
                translated_segments.append(segment)
        
        return "\n\n".join(translated_segments)
    
    except ImportError as e:
        error_log(f"[file] LLM 模組未安裝或無法導入：{e}")
        raise ImportError(f"需要 LLM 模組進行翻譯：{e}")


def _segment_text(text: str, max_length: int = 8000) -> list:
    """分段文本（避免單次處理過長）"""
    if len(text) <= max_length:
        return [text]
    
    # 簡單分段：按段落或固定長度
    info_log("[file] 文本過長，進行分段處理")
    segments = []
    
    # 優先按雙換行符分段（段落）
    paragraphs = text.split('\n\n')
    current_segment = ""
    
    for para in paragraphs:
        if len(current_segment) + len(para) <= max_length:
            current_segment += para + "\n\n"
        else:
            if current_segment:
                segments.append(current_segment.strip())
            
            # 如果單個段落太長，按單換行符分段
            if len(para) > max_length:
                lines = para.split('\n')
                temp_segment = ""
                for line in lines:
                    if len(temp_segment) + len(line) <= max_length:
                        temp_segment += line + "\n"
                    else:
                        if temp_segment:
                            segments.append(temp_segment.strip())
                        temp_segment = line + "\n"
                if temp_segment:
                    current_segment = temp_segment
                else:
                    current_segment = ""
            else:
                current_segment = para + "\n\n"
    
    if current_segment:
        segments.append(current_segment.strip())
    
    return segments


def _save_translated_file(text: str, output_path: str, file_ext: str):
    """儲存翻譯後的文件"""
    output_path_obj = Path(output_path)
    output_path_obj.parent.mkdir(parents=True, exist_ok=True)
    
    if file_ext == ".txt" or file_ext == ".md":
        output_path_obj.write_text(text, encoding="utf-8")
    
    elif file_ext == ".pdf":
        # PDF 轉存為 TXT（因為 PDF 生成較複雜）
        txt_path = output_path_obj.with_suffix(".txt")
        txt_path.write_text(text, encoding="utf-8")
        info_log(f"[file] PDF 翻譯結果已儲存為 TXT：{txt_path}")
        return str(txt_path)
    
    elif file_ext == ".docx":
        try:
            from docx import Document
            doc = Document()
            for para in text.split('\n\n'):
                doc.add_paragraph(para)
            doc.save(output_path)
        except ImportError:
            # 降級為 TXT
            txt_path = output_path_obj.with_suffix(".txt")
            txt_path.write_text(text, encoding="utf-8")
            info_log(f"[file] python-docx 未安裝，翻譯結果已儲存為 TXT：{txt_path}")
            return str(txt_path)
    
